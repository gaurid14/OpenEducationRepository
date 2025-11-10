import asyncio
import threading
import urllib

import docx
from asgiref.sync import async_to_sync
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from docx import Document

from accounts.models import Chapter, UploadCheck, Assessment, Question, Option, Course
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseUpload, MediaIoBaseDownload
from google.oauth2.credentials import Credentials
from django.conf import settings
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse, HttpResponseBadRequest
import os
import tempfile
import json
import io
from urllib.parse import unquote

import pdfkit
import time

from langgraph_agents.agents.submission_agent import submission_agent
from langgraph_agents.graph.workflow import compiled_graph, graph
from langgraph_agents.services.drive_service import get_drive_service, get_or_create_drive_folder
from langgraph_agents.services.gemini_service import llm

from urllib.parse import unquote


def contributor_upload_file(request):
    course_id = request.GET.get('course_id')
    chapter_id = request.GET.get('chapter_id')
    topic = unquote(request.GET.get('topic', ''))

    service = get_drive_service()

    # Validate inputs
    if not all([course_id, chapter_id, topic]):
        return HttpResponseBadRequest("Missing course_id, chapter_id, or topic parameter")

    # Fetch course and chapter safely
    course = get_object_or_404(Course, id=course_id)
    chapter = get_object_or_404(Chapter, id=chapter_id)
    contributor_id = request.user.id

    # 🔹 Check if already submitted
    existing_upload = UploadCheck.objects.filter(
        contributor_id=contributor_id,
        chapter_id=chapter_id
    ).first()

    if existing_upload:
        return render(request, "contributor/after_submission.html", {
            "chapter_name": chapter.chapter_name,
            "contributor_id": contributor_id,
            "message": "You have already submitted this chapter’s content."
        })

    # Store context in session
    request.session.update({
        "contributor_id": contributor_id,
        "course_name": course.course_name,
        "course_id": course_id,
        "chapter_id": chapter_id,
        "chapter_name": chapter.chapter_name,
        "chapter_number": chapter.chapter_number,
        "description": chapter.description,
        "topic": topic,
    })

    # Initialize Drive service
    service = get_drive_service()
    oer_root_id = get_or_create_drive_folder(service, "oer_content")

    # Base folder name for contributor-chapter
    base_folder_name = f"{contributor_id}_{course.id}_{chapter.chapter_number}"

    # Helper: fetch files for a folder type and topic
    def get_files_from_folder(folder_type):
        try:
            # 1️⃣ Find or create root subfolder (drafts/pdf/videos) under main OER folder
            root_folder_id = get_or_create_drive_folder(
                service, settings.GOOGLE_DRIVE_FOLDERS[folder_type], oer_root_id
            )

            # 2️⃣ Find contributor's folder inside that section
            query = (
                f"mimeType='application/vnd.google-apps.folder' "
                f"and name='{base_folder_name}' "
                f"and '{root_folder_id}' in parents and trashed=false"
            )
            folders = service.files().list(q=query, fields="files(id, name)").execute().get('files', [])
            if not folders:
                return []  # Contributor folder doesn’t exist

            chapter_folder_id = folders[0]['id']

            # 3️⃣ Now check for topic folder inside contributor’s folder
            query_topic = (
                f"mimeType='application/vnd.google-apps.folder' "
                f"and name='{topic}' "
                f"and '{chapter_folder_id}' in parents and trashed=false"
            )
            topic_folders = service.files().list(q=query_topic, fields="files(id, name)").execute().get('files', [])
            if not topic_folders:
                return []  # Topic folder doesn’t exist yet

            topic_folder_id = topic_folders[0]['id']

            # 4️⃣ Fetch all files under that topic folder
            result = service.files().list(
                q=f"'{topic_folder_id}' in parents and trashed=false",
                fields="files(id, name, mimeType)"
            ).execute()

            files_list = []
            for f in result.get('files', []):
                files_list.append({
                    'id': f['id'],
                    'name': f['name'],
                    'mimeType': f['mimeType'],
                    'type': folder_type
                })
            return files_list

        except Exception as e:
            print(f"[Drive Fetch Error - {folder_type}] {e}")
            return []

    # Fetch all file types
    files = []
    for folder_type in ['drafts', 'pdf', 'videos']:
        files.extend(get_files_from_folder(folder_type))

    context = {
        "course": course,
        "chapter": chapter,
        "topic": topic,
        "files": files,
    }
    return render(request, "contributor/contributor_upload_file.html", context)


@csrf_exempt
def upload_files(request):
    """Upload PDFs or videos to Drive — organized by contributor, chapter, and topic (keep topic name intact)."""
    contributor_id = request.session.get('contributor_id')
    course_id = request.session.get('course_id')
    chapter_number = request.session.get('chapter_number')
    chapter_name = request.session.get('chapter_name')
    topic_name = (request.POST.get('topic') or request.GET.get('topic') or "").strip()
    print(f"[UPLOAD] Chapter: {chapter_name}, Topic: '{topic_name}'")

    if request.method != "POST":
        return HttpResponseBadRequest("Invalid request")

    service = get_drive_service()
    oer_root_id = get_or_create_drive_folder(service, "oer_content")

    base_folder_name = f"{contributor_id}_{course_id}_{chapter_number}"
    safe_topic_name = topic_name.replace("/", "-").replace("\\", "-")

    # ✅ Read all files first before uploading
    pdf_files = request.FILES.getlist('pdf_file')
    video_files = request.FILES.getlist('video_file')

    def get_or_create_nested_folder(service, parent_id, folder_name):
        query = (
            f"mimeType='application/vnd.google-apps.folder' "
            f"and name='{folder_name}' "
            f"and '{parent_id}' in parents and trashed=false"
        )
        results = service.files().list(q=query, fields="files(id, name)").execute()
        folders = results.get("files", [])
        if folders:
            return folders[0]["id"]
        folder_metadata = {
            "name": folder_name,
            "mimeType": "application/vnd.google-apps.folder",
            "parents": [parent_id],
        }
        new_folder = service.files().create(body=folder_metadata, fields="id").execute()
        return new_folder["id"]

    def ensure_topic_folder(folder_type):
        type_root = get_or_create_drive_folder(service, settings.GOOGLE_DRIVE_FOLDERS[folder_type], oer_root_id)
        contrib_folder = get_or_create_nested_folder(service, type_root, base_folder_name)
        if safe_topic_name:
            topic_folder = get_or_create_nested_folder(service, contrib_folder, safe_topic_name)
            return topic_folder
        return contrib_folder

    # =================== PDF UPLOAD ===================
    if pdf_files:
        topic_pdf_folder = ensure_topic_folder("pdf")
        print(f"[UPLOAD] Uploading {len(pdf_files)} PDFs to folder ID {topic_pdf_folder}")
        for pdf in pdf_files:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                for chunk in pdf.chunks():
                    tmp.write(chunk)
                tmp.flush()
                tmp_path = tmp.name

            media = MediaFileUpload(tmp_path, mimetype="application/pdf", resumable=False)
            service.files().create(
                body={'name': f"{contributor_id}_{pdf.name}", 'parents': [topic_pdf_folder]},
                media_body=media,
                fields='id'
            ).execute()

            try:
                media._fd.close()
                os.remove(tmp_path)
            except Exception as e:
                print(f"[ERROR] Could not delete temp file: {tmp_path}, {e}")

    # =================== VIDEO UPLOAD ===================
    if video_files:
        topic_video_folder = ensure_topic_folder("videos")
        print(f"[UPLOAD] Uploading {len(video_files)} videos to folder ID {topic_video_folder}")
        for video in video_files:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as tmp:
                for chunk in video.chunks():
                    tmp.write(chunk)
                tmp.flush()
                tmp_path = tmp.name

            media = MediaFileUpload(tmp_path, mimetype="video/mp4", resumable=False)
            service.files().create(
                body={'name': f"{contributor_id}_{video.name}", 'parents': [topic_video_folder]},
                media_body=media,
                fields='id'
            ).execute()

            try:
                media._fd.close()
                os.remove(tmp_path)
            except Exception as e:
                print(f"[ERROR] Could not delete temp file: {tmp_path}, {e}")

    # ✅ Redirect only after both are done
    messages.success(request, "Files uploaded to Google Drive successfully!")

    course_id = request.session.get("course_id")
    chapter_id = request.session.get("chapter_id")
    request.GET = request.GET.copy()
    request.GET["course_id"] = str(course_id)
    request.GET["chapter_id"] = str(chapter_id)
    request.GET["topic"] = topic_name or ""

    return redirect(
        f"/dashboard/contributor/submit_content/upload/submission?"
        f"course_id={course_id}&chapter_id={chapter_id}&topic={urllib.parse.quote(topic_name)}"
    )

    # return contributor_upload_file(request)






    # return render(request, "contributor/submit_content.html")


# ---------------- EDITOR / DRAFT ---------------- #
@csrf_exempt
def contributor_editor(request):
    """Save drafts as HTML or final submissions as PDF in Google Drive."""
    service = get_drive_service()

    contributor_id = request.session.get('contributor_id', 101)
    course_id = request.session.get('course_id')
    chapter_number = request.session.get('chapter_number')
    topic_name = request.POST.get('topic') or request.GET.get('topic')
    print("Topic name: " + topic_name)
    chapter_name = request.session.get('chapter_name', 'structured_query_language')

    # Root folder for all OER content
    oer_root_id = get_or_create_drive_folder(service, "oer_content")
    drafts_root_id = get_or_create_drive_folder(service, settings.GOOGLE_DRIVE_FOLDERS['drafts'], oer_root_id)
    pdf_root_id = get_or_create_drive_folder(service, settings.GOOGLE_DRIVE_FOLDERS['pdf'], oer_root_id)

    # Contributor-specific folders
    # drafts_folder_id = get_or_create_drive_folder(service, f"{contributor_id}_{course_id}_{chapter_number}", drafts_root_id)
    # pdf_folder_id = get_or_create_drive_folder(service, f"{contributor_id}_{course_id}_{chapter_number}", pdf_root_id)

    # Contributor-level folders
    base_folder_name = f"{contributor_id}_{course_id}_{chapter_number}"
    drafts_folder_id = get_or_create_drive_folder(service, base_folder_name, drafts_root_id)
    pdf_folder_id = get_or_create_drive_folder(service, base_folder_name, pdf_root_id)

    # ✅ Topic-level subfolders inside contributor folder
    if topic_name:
        safe_topic_name = topic_name.replace("/", "_").strip()  # avoid invalid path characters
        drafts_topic_folder_id = get_or_create_drive_folder(service, safe_topic_name, drafts_folder_id)
        pdf_topic_folder_id = get_or_create_drive_folder(service, safe_topic_name, pdf_folder_id)
    else:
        drafts_topic_folder_id = drafts_folder_id
        pdf_topic_folder_id = pdf_folder_id

    if request.method == 'POST':
        action = request.POST.get('action')  # 'draft' or 'submitDraft'
        content = request.POST.get('notes', '')
        filename = request.POST.get('filename', 'draft.html')
        file_id = request.POST.get('file_id')  # New hidden input in form

        try:
            if action == 'draft':
                # Convert HTML content to DOCX
                doc = Document()
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(content, 'html.parser')
                for paragraph in soup.find_all(['p', 'div']):
                    text = paragraph.get_text(strip=True)
                    if text:
                        doc.add_paragraph(text)

                file_io = io.BytesIO()
                doc.save(file_io)
                file_io.seek(0)
                media = MediaIoBaseUpload(
                    file_io,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    resumable=True
                )

                user_filename = filename if filename.lower().endswith('.docx') else f"{filename}.docx"
                doc_filename = f"{contributor_id}_{user_filename}"

                if file_id:
                    # Update existing draft
                    service.files().update(
                        fileId=file_id,
                        media_body=media
                    ).execute()
                else:
                    # Create new draft
                    doc_filename = f"{contributor_id}_{user_filename}"
                    service.files().create(
                        body={'name': doc_filename, 'parents': [drafts_topic_folder_id]},
                        media_body=media,
                        fields='id'
                    ).execute()

            elif action == 'submitDraft':
                # Save as PDF using xhtml2pdf
                from xhtml2pdf import pisa

                pdf_io = io.BytesIO()
                # Convert HTML to PDF
                result = pisa.CreatePDF(io.StringIO(content), dest=pdf_io)
                if result.err:
                    raise Exception("PDF generation failed")

                pdf_io.seek(0)
                media = MediaIoBaseUpload(pdf_io, mimetype='application/pdf', resumable=True)

                # Use same naming as draft
                # Ensure filename ends with .pdf
                user_filename = filename if filename.lower().endswith('.pdf') else f"{filename}.pdf"

                # Prepend contributor_id only if not already present
                if not user_filename.startswith(f"{contributor_id}_"):
                    pdf_filename = f"{contributor_id}_{user_filename}"
                else:
                    pdf_filename = user_filename

                # Upload PDF to Drive
                service.files().create(
                    body={'name': pdf_filename, 'parents': [pdf_topic_folder_id]},
                    media_body=media,
                    fields='id'
                ).execute()

                # Delete the draft from Drive if editing an existing draft
                if file_id:
                    try:
                        service.files().delete(fileId=file_id).execute()
                    except Exception as e:
                        print(f"[WARNING] Could not delete draft {file_id}: {e}")

        except Exception as e:
            print(f"[ERROR] {action.capitalize()} upload failed: {e}")
            messages.error(request, f"Failed to save {action}: {e}")

    # Fetch existing drafts
    try:
        query = f"'{drafts_folder_id}' in parents and trashed=false"
        results = service.files().list(q=query, fields="files(id, name, createdTime)").execute()
        files = results.get('files', [])
    except Exception as e:
        files = []
        print(f"[ERROR] Failed to fetch drafts: {e}")

    course_id = request.session.get("course_id")
    chapter_id = request.session.get("chapter_id")

    request.GET = request.GET.copy()
    request.GET['course_id'] = str(course_id)
    request.GET['chapter_id'] = str(chapter_id)
    request.GET['topic'] = topic_name or ''

    return contributor_upload_file(request)


# ---------------- LOAD FILE CONTENT ---------------- #
@csrf_exempt
def load_file(request):
    service = get_drive_service()
    file_id = request.GET.get('file_id')

    if not file_id:
        return JsonResponse({'error': 'file_id is required'}, status=400)

    try:
        request_file = service.files().get(fileId=file_id, fields='mimeType').execute()
        mime_type = request_file.get('mimeType', 'text/html')

        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, service.files().get_media(fileId=file_id))
        done = False
        while not done:
            _, done = downloader.next_chunk()
        fh.seek(0)

        print("Loading file_id:", file_id)
        print("Service object:", service)

        if 'text/html' in mime_type:
            content = fh.getvalue().decode('utf-8')
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            import docx
            doc = docx.Document(fh)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Join with <p> tags for TinyMCE
            content = ''.join(f'<p>{p}</p>' for p in paragraphs)
        else:
            content = f"<p>Cannot edit file of type {mime_type} in the editor.</p>"

        return JsonResponse({'content': content})

    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
def delete_drive_file(request):
    """Delete a Google Drive file permanently."""
    if request.method == 'POST':
        file_id = request.POST.get('file_id')
        service = get_drive_service()
        try:
            service.files().delete(fileId=file_id).execute()
            return JsonResponse({'success': True, 'message': 'File deleted successfully.'})
        except Exception as e:
            return JsonResponse({'success': False, 'message': str(e)})
    # Redirect back to original submission page
    course_id = request.session.get("course_id")
    chapter_id = request.session.get("chapter_id")
    return redirect(f'/dashboard/contributor/submit_content/?course_id={course_id}&chapter_id={chapter_id}')
    # return JsonResponse({'success': False, 'message': 'Invalid request'})


@csrf_exempt
def submit_assessment(request):
    course_id = request.session.get("course_id")
    chapter_id = request.session.get("chapter_id")

    if not course_id or not chapter_id:
        messages.error(request, "Course or Chapter not found in session.")
        return redirect("/dashboard/contributor/submit_content/")

    course = Course.objects.get(id=course_id)
    chapter = Chapter.objects.get(id=chapter_id)

    if request.method == 'POST':
        # Extract all questions dynamically
        questions_data = []
        i = 0
        while f'questions[{i}][question]' in request.POST:
            q_text = request.POST[f'questions[{i}][question]']
            correct = int(request.POST[f'questions[{i}][correct]'])
            options = request.POST.getlist(f'questions[{i}][options][]')
            questions_data.append({'text': q_text, 'correct': correct, 'options': options})
            i += 1

        assessment = Assessment.objects.create(course=course, chapter=chapter, contributor_id=request.user)

        for q in questions_data:
            question = Question.objects.create(
                assessment=assessment,
                text=q['text'],
                correct_option=q['correct']
            )
            for opt_text in q['options']:
                Option.objects.create(question=question, text=opt_text)

    return redirect(f'/dashboard/contributor/submit_content/?course_id={course_id}&chapter_id={chapter_id}')


@csrf_exempt
def gemini_chat(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            user_message = data.get('message', '')
            if not user_message:
                return JsonResponse({'error': 'No message provided'}, status=400)

            # Send to Gemini API
            response = llm.predict(user_message)

            return JsonResponse({'reply': response})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Invalid method'}, status=405)


@csrf_exempt
def confirm_submission(request):
    """Handles final submission confirmation and triggers the LangGraph submission agent."""
    if request.method != "POST":
        return JsonResponse({'error': 'POST required'}, status=405)

    try:
        contributor_id = request.session.get('contributor_id')
        course_id = request.session.get('course_id')
        chapter_id = request.session.get('chapter_id')

        if not all([contributor_id, course_id, chapter_id]):
            return JsonResponse({'error': 'Missing session data'}, status=400)

        chapter_obj = Chapter.objects.get(id=chapter_id)
        chapter_name = chapter_obj.chapter_name
        chapter_number = chapter_obj.chapter_number

        service = get_drive_service()
        oer_root_id = get_or_create_drive_folder(service, "oer_content")

        # 🔹 Create or get subfolders for each content type
        pdf_root_id = get_or_create_drive_folder(service, settings.GOOGLE_DRIVE_FOLDERS['pdf'], oer_root_id)
        video_root_id = get_or_create_drive_folder(service, settings.GOOGLE_DRIVE_FOLDERS['videos'], oer_root_id)
        assess_root_id = get_or_create_drive_folder(service, settings.GOOGLE_DRIVE_FOLDERS['assessments'], oer_root_id)

        # 🔹 Contributor-specific folders
        pdf_folder_id = get_or_create_drive_folder(service, f"{contributor_id}_{course_id}_{chapter_number}",
                                                   pdf_root_id)
        video_folder_id = get_or_create_drive_folder(service, f"{contributor_id}_{course_id}_{chapter_number}",
                                                     video_root_id)
        assess_folder_id = get_or_create_drive_folder(service, f"{contributor_id}_{course_id}_{chapter_number}",
                                                      assess_root_id)

        # 🔹 Prepare LangGraph state
        state = {
            "contributor_id": contributor_id,
            "chapter_name": chapter_name,
            "chapter_id": chapter_id,
            "course_id": course_id,
            "drive_folders": {
                "pdf": pdf_folder_id,
                "videos": video_folder_id,
                "assessments": assess_folder_id
            }
        }

        # async def run_graph():
        #     return await compiled_graph.ainvoke(state)       # compiled_graph is awaitable
        #
        # # Run synchronously (not background)
        # result = asyncio.run(run_graph())
        #
        # # 🔹 Check result before redirecting
        # if result and result.get("status") == "submission_recorded":
        #     return render(request, "contributor/after_submission.html", {
        #         "chapter_name": chapter_name,
        #         "contributor_id": contributor_id
        #     })

        # Run submission agent synchronously
        async def run_submission():
            return await submission_agent.ainvoke(state)

        result = asyncio.run(run_submission())  # UI waits only for DB update

        if result.get("status") == "submission_recorded":
            # Trigger evaluation agent in background (non-blocking)
            import threading
            def run_graph_in_thread(state):
                asyncio.run(compiled_graph.ainvoke(state))  # ensures coroutine is executed

            # trigger in a thread
            threading.Thread(target=run_graph_in_thread, args=(state,)).start()

            return render(request, "contributor/final_submission.html", {
                "chapter_name": chapter_name,
                "contributor_id": contributor_id
            })

        else:
            print(f"[ERROR] Submission agent did not return success: {result}")
            return JsonResponse({'error': 'Submission agent failed or returned invalid response'}, status=500)

    except Exception as e:
        print(f"[ERROR] Final submission failed: {e}")
        return JsonResponse({'error': str(e)}, status=500)


# # Inside the upload_files view function...
# def upload_files(request):
#     # Get course_id, chapter_id, topic_name needed for the redirect
#     # These might come from hidden inputs in the form or session
#     course_id = request.POST.get('course_id')
#     chapter_id = request.POST.get('chapter_id')
#     topic_name = request.POST.get('topic_name', 'Unknown Topic') # Get topic name
#
#     if request.method == "POST":
#         # ... your logic to handle file saving to Drive ...
#
#         messages.success(request, "Files uploaded successfully!")
#
#         # FIX: Redirect back to the SAME upload page
#         # Make sure your 'contributor_upload_file' URL takes these parameters
#         return redirect()
#
#     # Handle GET request (if this view also displays the initial form)
#     # return render(request, 'accounts/contributor_submit_content.html', ...)
def generate_assessment(request):
    print("Assessment view called")
    # generate_expertise()
    # Clear all session data safely
    return render(request, 'contributor/generated-assessment.html')


def after_submission(request):
    print("After submission view called")
    # generate_expertise()
    # Clear all session data safely
    return render(request, 'contributor/after_submission.html')


def final_submission(request):
    print("Final view called")
    # generate_expertise()
    # Clear all session data safely
    return render(request, 'contributor/final_submission.html')
